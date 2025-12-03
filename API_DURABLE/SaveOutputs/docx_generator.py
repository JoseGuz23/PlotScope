# =============================================================================
# docx_generator.py - SYLPHRENA 4.0
# =============================================================================
# Genera documento DOCX profesional con Track Changes simulados
# Compatible con Microsoft Word y Google Docs
# =============================================================================

import logging
from datetime import datetime
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import re

logging.basicConfig(level=logging.INFO)


def generate_docx_with_track_changes(chapters: list, book_name: str) -> BytesIO:
    """
    Genera documento DOCX con cambios marcados profesionalmente.
    
    Args:
        chapters: Lista de capítulos consolidados con contenido_original, 
                  contenido_editado y cambios_realizados
        book_name: Nombre del libro
    
    Returns:
        BytesIO: Documento DOCX en memoria listo para guardar
    """
    logging.info(f"📄 Generando DOCX con Track Changes: {book_name}")
    
    doc = Document()
    
    # =============================================================================
    # PORTADA
    # =============================================================================
    add_title_page(doc, book_name)
    
    # =============================================================================
    # CAPÍTULOS CON CAMBIOS MARCADOS
    # =============================================================================
    for i, chapter in enumerate(chapters):
        logging.info(f"   Procesando: {chapter.get('display_title', f'Cap {i+1}')}")
        add_chapter_with_changes(doc, chapter)
    
    # =============================================================================
    # GUARDAR EN MEMORIA
    # =============================================================================
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    logging.info(f"✅ DOCX generado: {len(chapters)} capítulos")
    return buffer


def add_title_page(doc: Document, book_name: str):
    """Agrega portada profesional."""
    # Título principal
    title = doc.add_heading(book_name.upper(), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtítulo
    subtitle = doc.add_paragraph('Manuscrito Editado')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(16)
    subtitle_format.font.color.rgb = RGBColor(102, 102, 102)
    
    # Info de edición
    doc.add_paragraph()  # Espacio
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.add_run(f'Editado por Sylphrena 4.0\n{datetime.now().strftime("%d de %B, %Y")}')
    info_run.font.size = Pt(12)
    info_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Instrucciones
    doc.add_paragraph()
    doc.add_paragraph()
    instructions = doc.add_paragraph()
    instructions.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inst_run = instructions.add_run(
        '═══════════════════════════════════════\n'
        'INSTRUCCIONES\n'
        '═══════════════════════════════════════'
    )
    inst_run.font.size = Pt(10)
    inst_run.bold = True
    
    doc.add_paragraph()
    guide = doc.add_paragraph()
    guide_text = guide.add_run(
        '🔴 Texto tachado en rojo = ELIMINADO\n'
        '🟢 Texto en verde = AÑADIDO\n'
        '💬 Comentarios al margen = Justificación de cambios\n\n'
        'Puede aceptar o rechazar cambios usando:\n'
        '• Microsoft Word: Pestaña "Revisar"\n'
        '• Google Docs: Herramientas → "Revisar sugerencias"'
    )
    guide_text.font.size = Pt(10)
    
    # Salto de página
    doc.add_page_break()


def add_chapter_with_changes(doc: Document, chapter: dict):
    """
    Agrega capítulo con cambios marcados inline.
    
    Estrategia:
    1. Identificar posiciones de cambios en texto original
    2. Reconstruir texto insertando marcas de Track Changes
    3. Agregar comentarios con justificaciones
    """
    display_title = chapter.get('display_title', 'Sin título')
    original = chapter.get('contenido_original', '')
    edited = chapter.get('contenido_editado', '')
    cambios = chapter.get('cambios_realizados', [])
    notas = chapter.get('notas_editor', '')
    
    # Título del capítulo
    doc.add_heading(display_title, 1)
    
    # Nota del editor si existe
    if notas:
        add_editor_note(doc, notas, len(cambios))
    
    # Si no hay cambios, insertar texto limpio
    if not cambios or original == edited:
        doc.add_paragraph(edited if edited else original)
        doc.add_page_break()
        return
    
    # Reconstruir texto con cambios marcados
    reconstruct_text_with_changes(doc, original, edited, cambios)
    
    # Salto de página
    doc.add_page_break()


def add_editor_note(doc: Document, notas: str, num_cambios: int):
    """Agrega nota del editor destacada."""
    note_para = doc.add_paragraph()
    note_para.paragraph_format.left_indent = Inches(0.5)
    note_para.paragraph_format.right_indent = Inches(0.5)
    
    # Icono y header
    header = note_para.add_run(f'📝 NOTA DEL EDITOR ({num_cambios} cambios)\n')
    header.bold = True
    header.font.size = Pt(10)
    header.font.color.rgb = RGBColor(0, 102, 204)
    
    # Contenido
    content = note_para.add_run(notas)
    content.font.size = Pt(9)
    content.font.italic = True
    content.font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph()  # Espacio


def reconstruct_text_with_changes(doc: Document, original: str, edited: str, cambios: list):
    """
    Reconstruye el texto marcando cambios inline.
    
    Estrategia simplificada:
    - Busca cada fragmento 'original' del cambio en el texto
    - Lo marca como eliminado (tachado rojo)
    - Inserta el 'editado' en verde justo después
    - Agrega comentario con justificación
    """
    # Empezar con el texto editado (resultado final)
    texto_actual = edited
    
    # Crear mapa de posiciones de cambios
    cambios_mapeados = []
    
    for i, cambio in enumerate(cambios):
        original_frag = cambio.get('original', '').strip()
        editado_frag = cambio.get('editado', '').strip()
        justificacion = cambio.get('justificacion', 'Sin justificación')
        tipo = cambio.get('tipo', 'otro').upper()
        
        if not original_frag:
            continue
        
        # Buscar fragmento en texto editado (o sus primeras palabras)
        # Esto es heurístico - buscamos contexto
        palabras_clave = original_frag.split()[:5]  # Primeras 5 palabras
        busqueda = ' '.join(palabras_clave)
        
        cambios_mapeados.append({
            'original': original_frag,
            'editado': editado_frag,
            'tipo': tipo,
            'justificacion': justificacion,
            'busqueda': busqueda
        })
    
    # Dividir texto en párrafos
    parrafos = texto_actual.split('\n\n')
    
    for parrafo_texto in parrafos:
        if not parrafo_texto.strip():
            continue
        
        para = doc.add_paragraph()
        
        # Buscar si este párrafo contiene algún cambio
        cambio_encontrado = None
        for cambio in cambios_mapeados:
            if cambio['busqueda'] in parrafo_texto or cambio['editado'] in parrafo_texto:
                cambio_encontrado = cambio
                break
        
        if cambio_encontrado:
            # Este párrafo tiene cambios - marcarlo
            add_paragraph_with_inline_changes(
                para, 
                parrafo_texto, 
                cambio_encontrado['original'],
                cambio_encontrado['editado'],
                cambio_encontrado['tipo'],
                cambio_encontrado['justificacion']
            )
        else:
            # Párrafo sin cambios
            para.add_run(parrafo_texto)


def add_paragraph_with_inline_changes(para, texto: str, original: str, editado: str, tipo: str, justificacion: str):
    """
    Agrega párrafo con cambios marcados inline.
    
    Formato:
    - Texto antes del cambio (normal)
    - [TEXTO ELIMINADO] (rojo tachado)
    - [TEXTO AÑADIDO] (verde)
    - Texto después del cambio (normal)
    """
    # Intentar encontrar dónde está el cambio
    if editado in texto:
        # El texto editado está en el párrafo
        partes = texto.split(editado, 1)
        
        # Antes del cambio
        if partes[0]:
            para.add_run(partes[0])
        
        # Mostrar eliminación
        if original:
            deleted_run = para.add_run(f"[{original}]")
            deleted_run.font.strike = True
            deleted_run.font.color.rgb = RGBColor(220, 53, 69)
            deleted_run.font.size = Pt(11)
        
        # Mostrar adición
        added_run = para.add_run(editado)
        added_run.font.color.rgb = RGBColor(40, 167, 69)
        added_run.font.bold = True
        added_run.font.size = Pt(11)
        
        # Agregar comentario visual inline
        comment_run = para.add_run(f" 💬")
        comment_run.font.color.rgb = RGBColor(0, 123, 255)
        comment_run.font.size = Pt(8)
        
        # Después del cambio
        if len(partes) > 1 and partes[1]:
            para.add_run(partes[1])
        
        # Agregar justificación como párrafo pequeño
        add_inline_comment(para.insert_paragraph_before(), tipo, justificacion)
    else:
        # No se encontró el cambio exacto - mostrar párrafo normal
        para.add_run(texto)


def add_inline_comment(para, tipo: str, justificacion: str):
    """Agrega comentario inline estilo Word."""
    para.paragraph_format.left_indent = Inches(0.8)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    
    # Tipo de cambio
    tipo_run = para.add_run(f"[{tipo}] ")
    tipo_run.font.size = Pt(8)
    tipo_run.bold = True
    tipo_run.font.color.rgb = RGBColor(0, 123, 255)
    
    # Justificación
    just_run = para.add_run(justificacion)
    just_run.font.size = Pt(8)
    just_run.font.italic = True
    just_run.font.color.rgb = RGBColor(108, 117, 125)


def create_simple_docx_with_changes(chapters: list, book_name: str) -> BytesIO:
    """
    Versión simplificada y más robusta.
    
    En lugar de intentar mapear cambios complejos:
    1. Muestra el texto editado completo
    2. Agrega tabla de cambios al final de cada capítulo
    """
    logging.info(f"📄 Generando DOCX simple con tabla de cambios: {book_name}")
    
    doc = Document()
    
    # Portada
    add_title_page(doc, book_name)
    
    # Capítulos
    for chapter in chapters:
        display_title = chapter.get('display_title', 'Sin título')
        edited = chapter.get('contenido_editado', '')
        cambios = chapter.get('cambios_realizados', [])
        notas = chapter.get('notas_editor', '')
        
        # Título
        doc.add_heading(display_title, 1)
        
        # Nota del editor
        if notas:
            add_editor_note(doc, notas, len(cambios))
        
        # Texto editado
        for parrafo in edited.split('\n\n'):
            if parrafo.strip():
                doc.add_paragraph(parrafo.strip())
        
        # Tabla de cambios
        if cambios:
            doc.add_paragraph()
            doc.add_heading('📋 Cambios Realizados', 2)
            add_changes_table(doc, cambios)
        
        doc.add_page_break()
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    logging.info(f"✅ DOCX simple generado")
    return buffer


def add_changes_table(doc: Document, cambios: list):
    """Agrega tabla profesional de cambios."""
    # Crear tabla
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Headers
    headers = table.rows[0].cells
    headers[0].text = 'Tipo'
    headers[1].text = 'Original'
    headers[2].text = 'Editado'
    headers[3].text = 'Justificación'
    
    # Formatear headers
    for cell in headers:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    
    # Agregar cambios
    for cambio in cambios:
        row = table.add_row().cells
        
        tipo = cambio.get('tipo', 'otro').upper()
        original = cambio.get('original', '')[:100] + ('...' if len(cambio.get('original', '')) > 100 else '')
        editado = cambio.get('editado', '')[:100] + ('...' if len(cambio.get('editado', '')) > 100 else '')
        justif = cambio.get('justificacion', '')
        
        row[0].text = tipo
        row[1].text = original
        row[2].text = editado
        row[3].text = justif
        
        # Formatear
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)


# =============================================================================
# FUNCIÓN PRINCIPAL EXPORTABLE
# =============================================================================

def generate_manuscript_docx(chapters: list, book_name: str, style: str = 'simple') -> BytesIO:
    """
    Genera manuscrito DOCX profesional.
    
    Args:
        chapters: Lista de capítulos consolidados
        book_name: Nombre del libro
        style: 'simple' (tabla de cambios) o 'inline' (cambios en texto)
    
    Returns:
        BytesIO con el documento
    """
    if style == 'inline':
        return generate_docx_with_track_changes(chapters, book_name)
    else:
        return create_simple_docx_with_changes(chapters, book_name)