# =============================================================================
# SegmentBook/__init__.py - SYLPHRENA 4.1.0 (BLOB STORAGE)
# =============================================================================
# MEJORAS:
#   - Lee archivos desde Azure Blob Storage
#   - Fallback a archivo local para desarrollo
#   - Soporta input como string (blob_path) o dict con configuración
# =============================================================================

import azure.functions as func
import regex as re
import json
import logging
import os
import sys
import traceback
from io import BytesIO

# Azure Storage
try:
    from azure.storage.blob import BlobServiceClient
    BLOB_AVAILABLE = True
except ImportError:
    BLOB_AVAILABLE = False

# Importaciones opcionales para formatos
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

MAX_CHARS_PER_CHUNK = 12000
DEFAULT_LIMIT_CHAPTERS = None 
MIN_CONTENT_CHARS = 100

logging.basicConfig(level=logging.INFO)

# -----------------------------------------------------------------------------
# FUNCIONES DE LECTURA
# -----------------------------------------------------------------------------

def read_from_blob_storage(blob_path: str) -> bytes:
    """Lee un archivo desde Azure Blob Storage."""
    if not BLOB_AVAILABLE:
        raise ImportError("azure-storage-blob no está instalado")
    
    connection_string = os.environ.get('AzureWebJobsStorage')
    if not connection_string:
        raise ValueError("AzureWebJobsStorage no configurado")
    
    # Parsear blob_path: "job_id/filename.docx"
    parts = blob_path.split('/')
    if len(parts) < 2:
        raise ValueError(f"blob_path inválido: {blob_path}")
    
    container_name = "sylphrena-inputs"
    blob_name = blob_path
    
    logging.info(f"📥 Leyendo desde Blob Storage: {container_name}/{blob_name}")
    
    blob_service = BlobServiceClient.from_connection_string(connection_string)
    blob_client = blob_service.get_blob_client(container_name, blob_name)
    
    if not blob_client.exists():
        raise FileNotFoundError(f"Blob no encontrado: {container_name}/{blob_name}")
    
    return blob_client.download_blob().readall()


def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    """Extrae texto de bytes (para blob storage)."""
    extension = os.path.splitext(filename)[1].lower()
    
    if extension == '.docx':
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx no instalado")
        doc = Document(BytesIO(file_bytes))
        return "\n".join([para.text for para in doc.paragraphs])
    
    elif extension == '.pdf':
        if not PDF_AVAILABLE:
            raise ImportError("pdfplumber no instalado")
        text = ""
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text += (page.extract_text() or "") + "\n"
        return text
    
    elif extension == '.txt':
        return file_bytes.decode('utf-8')
    
    else:
        raise ValueError(f"Formato no soportado: {extension}")


def extract_text_from_file(file_path: str) -> str:
    """Extrae texto de archivo local (fallback para desarrollo)."""
    try:
        extension = os.path.splitext(file_path)[1].lower()
        
        if extension == '.pdf':
            if not PDF_AVAILABLE: raise ImportError("pdfplumber no instalado")
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += (page.extract_text() or "") + "\n"
            return text
        
        elif extension == '.docx':
            if not DOCX_AVAILABLE: raise ImportError("python-docx no instalado")
            doc = Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        
        elif extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        else:
            raise ValueError(f"Formato no soportado: {extension}")
            
    except Exception as e:
        logging.error(f"❌ Error extrayendo texto: {e}")
        raise


# -----------------------------------------------------------------------------
# FUNCIONES AUXILIARES
# -----------------------------------------------------------------------------

def detect_section_type(title: str) -> str:
    """Detecta el tipo de sección basado en el título."""
    title_lower = title.lower().strip()
    
    if any(kw in title_lower for kw in ['prólogo', 'prologo', 'prefacio']):
        return 'PROLOGUE'
    elif any(kw in title_lower for kw in ['epílogo', 'epilogo']):
        return 'EPILOGUE'
    elif any(kw in title_lower for kw in ['interludio', 'intermedio']):
        return 'INTERLUDE'
    elif any(kw in title_lower for kw in ['introducción', 'introduccion']):
        return 'INTRODUCTION'
    elif 'acto' in title_lower:
        return 'ACT_HEADER'
    elif 'parte' in title_lower:
        return 'PART_HEADER'
    elif 'capítulo' in title_lower or 'capitulo' in title_lower:
        return 'CHAPTER'
    elif 'nota' in title_lower and 'editor' in title_lower:
        return 'EDITOR_NOTE'
    else:
        return 'CHAPTER'


def smart_split(text: str, max_chars: int) -> list:
    """Divide texto largo en fragmentos respetando párrafos."""
    if len(text) <= max_chars: return [text]
    chunks = []
    while len(text) > max_chars:
        split_point = text.rfind('\n\n', 0, max_chars)
        if split_point == -1: split_point = text.rfind('\n', 0, max_chars)
        if split_point == -1: split_point = text.rfind('. ', 0, max_chars) + 1
        if split_point <= 0: split_point = max_chars
        
        chunks.append(text[:split_point].strip())
        text = text[split_point:].strip()
    if text: chunks.append(text)
    return chunks


def generate_hierarchical_metadata(chapters_raw: list) -> list:
    """Genera la estructura plana de fragmentos con metadatos."""
    final_list = []
    global_fragment_id = 1
    chapter_id = 1
    
    for chapter_data in chapters_raw:
        raw_title = chapter_data['title']
        content = chapter_data['content']
        section_type = chapter_data['section_type']
        
        if len(content) > MAX_CHARS_PER_CHUNK:
            sub_chunks = smart_split(content, MAX_CHARS_PER_CHUNK)
            total_frags = len(sub_chunks)
            for idx, chunk in enumerate(sub_chunks):
                final_list.append({
                    'id': global_fragment_id,
                    'parent_chapter_id': chapter_id,
                    'original_title': raw_title,
                    'title': f"{raw_title} ({idx + 1}/{total_frags})",
                    'fragment_index': idx + 1,
                    'total_fragments': total_frags,
                    'section_type': section_type,
                    'is_fragment': True,
                    'content': chunk,
                    'word_count': len(chunk.split())
                })
                global_fragment_id += 1
        else:
            final_list.append({
                'id': global_fragment_id,
                'parent_chapter_id': chapter_id,
                'original_title': raw_title,
                'title': raw_title,
                'fragment_index': 1,
                'total_fragments': 1,
                'section_type': section_type,
                'is_fragment': False,
                'content': content,
                'word_count': len(content.split())
            })
            global_fragment_id += 1
        chapter_id += 1
    return final_list


# -----------------------------------------------------------------------------
# MAIN FUNCTION (ACTIVITY TRIGGER)
# -----------------------------------------------------------------------------

def main(book_path) -> str:
    """
    Activity Function que recibe configuración y devuelve el libro segmentado.
    
    Input puede ser:
    - String: "job_id/filename.docx" (blob path)
    - Dict: { "blob_path": "...", "limit_chapters": 2 }
    """
    try:
        logging.info("=" * 80)
        logging.info("🚀 SegmentBook v4.1.0 (Blob Storage) iniciado")
        
        # --- PASO 1: PARSEAR INPUT ---
        blob_path = ""
        limit_chapters = DEFAULT_LIMIT_CHAPTERS
        
        if isinstance(book_path, dict):
            blob_path = book_path.get('blob_path', book_path.get('book_path', ''))
            if 'limit_chapters' in book_path:
                limit_chapters = book_path['limit_chapters']
                
        elif isinstance(book_path, str):
            if book_path.strip().startswith('{'):
                try:
                    data = json.loads(book_path)
                    blob_path = data.get('blob_path', data.get('book_path', ''))
                    if 'limit_chapters' in data:
                        limit_chapters = data['limit_chapters']
                except:
                    blob_path = book_path
            else:
                blob_path = book_path

        logging.info(f"📂 Input -> blob_path: {blob_path} | Límite: {limit_chapters}")
        
        # --- PASO 2: LEER ARCHIVO ---
        text = ""
        source_info = ""
        
        # Intentar leer desde Blob Storage
        if blob_path and '/' in blob_path and BLOB_AVAILABLE:
            try:
                filename = blob_path.split('/')[-1]
                file_bytes = read_from_blob_storage(blob_path)
                text = extract_text_from_bytes(file_bytes, filename)
                source_info = f"blob://{blob_path}"
                logging.info(f"✅ Archivo leído desde Blob Storage: {blob_path}")
            except Exception as e:
                logging.warning(f"⚠️ No se pudo leer de Blob Storage: {e}")
                logging.warning(f"   Intentando fallback a archivo local...")
        
        # Fallback: archivo local (para desarrollo)
        if not text:
            script_dir = os.path.dirname(os.path.abspath(__file__))
            
            # Buscar archivo por nombre si se proporcionó
            if blob_path:
                filename = blob_path.split('/')[-1] if '/' in blob_path else blob_path
                local_file_path = os.path.join(script_dir, filename)
            else:
                # Default para pruebas
                local_file_path = os.path.join(script_dir, 'Piel_Morena.docx')
            
            if os.path.exists(local_file_path):
                text = extract_text_from_file(local_file_path)
                source_info = f"local://{local_file_path}"
                logging.info(f"✅ Archivo leído localmente: {local_file_path}")
            else:
                # Buscar cualquier docx en el directorio
                for f in os.listdir(script_dir):
                    if f.endswith('.docx'):
                        local_file_path = os.path.join(script_dir, f)
                        text = extract_text_from_file(local_file_path)
                        source_info = f"local://{local_file_path}"
                        logging.info(f"✅ Usando archivo encontrado: {local_file_path}")
                        break
        
        if not text:
            raise FileNotFoundError(f"No se encontró archivo para procesar. blob_path={blob_path}")
        
        logging.info(f"📄 Texto extraído: {len(text):,} caracteres")
        
        # --- PASO 3: DETECTAR CAPÍTULOS ---
        special_keywords = r'(?:Prólogo|Prefacio|Introducción|Interludio|Epílogo|Nota para el editor)'
        full_pattern = f'(?mi)(?:^\\s*)(?:{special_keywords}|(?:Capítulo|Acto|Parte)\\s+)[^\n]*'
        original_chapters = re.split(f'(?={full_pattern})', text)
        
        chapters_raw = []
        skipped_headers = []
        
        for raw_chapter in original_chapters:
            if not raw_chapter.strip(): 
                continue
            
            lines = raw_chapter.strip().split('\n')
            title = lines[0].strip()
            content = '\n'.join(lines[1:]).strip()
            section_type = detect_section_type(title)
            content_length = len(content)
            
            is_structural_header = section_type in ('ACT_HEADER', 'PART_HEADER')
            
            if is_structural_header and content_length < MIN_CONTENT_CHARS:
                logging.warning(f"⚠️ Saltando encabezado estructural vacío: '{title[:50]}' ({content_length} chars)")
                skipped_headers.append({
                    'title': title,
                    'chars': content_length,
                    'type': section_type
                })
                continue
            
            if content_length < MIN_CONTENT_CHARS:
                logging.info(f"📝 Capítulo corto pero válido: '{title[:50]}' ({content_length} chars)")
            
            chapters_raw.append({
                'title': title,
                'content': content,
                'section_type': section_type
            })
            
            logging.info(f"📖 Capítulo detectado: '{title[:50]}' ({content_length:,} chars) - tipo: {section_type}")

        # --- PASO 4: APLICAR LÍMITE ---
        total_detected = len(chapters_raw)
        logging.info(f"📊 Total capítulos con contenido: {total_detected}")
        
        if limit_chapters is not None and isinstance(limit_chapters, int) and limit_chapters > 0:
            logging.warning(f"✂️ RECORTANDO: Procesando {limit_chapters} de {total_detected} capítulos válidos.")
            chapters_raw = chapters_raw[:limit_chapters]
        else:
            logging.info(f"✅ PROCESANDO LIBRO COMPLETO ({total_detected} capítulos).")

        # --- PASO 5: SEGMENTAR ---
        fragments = generate_hierarchical_metadata(chapters_raw) 

        chapter_map = {}
        for frag in fragments:
            p_id = frag['parent_chapter_id']
            if p_id not in chapter_map:
                chapter_map[p_id] = {'fragment_ids': [], 'original_title': frag['original_title']}
            chapter_map[p_id]['fragment_ids'].append(frag['id'])

        result = {
            'fragments': fragments,
            'book_metadata': {
                'total_chapters': len(chapter_map),
                'total_fragments': len(fragments),
                'skipped_headers': len(skipped_headers),
                'source': source_info
            },
            'chapter_map': chapter_map,
            'skipped_headers': skipped_headers
        }

        logging.info(f"")
        logging.info(f"{'='*60}")
        logging.info(f"✅ SEGMENTACIÓN COMPLETADA")
        logging.info(f"   Fuente: {source_info}")
        logging.info(f"   Capítulos válidos: {len(chapter_map)}")
        logging.info(f"   Fragmentos totales: {len(fragments)}")
        logging.info(f"   Encabezados saltados: {len(skipped_headers)}")
        logging.info(f"{'='*60}")
        
        return json.dumps(result, ensure_ascii=False)

    except Exception as e:
        logging.error(f"❌ Error Fatal en SegmentBook: {str(e)}")
        logging.error(traceback.format_exc())
        raise e